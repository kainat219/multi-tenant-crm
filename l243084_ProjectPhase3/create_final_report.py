from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "MultiTenant_CRM_Final_Project_Report.docx"


ACCENT = RGBColor(7, 89, 133)
GREEN = RGBColor(15, 118, 110)
GOLD = RGBColor(202, 138, 4)
DARK = RGBColor(23, 32, 51)
MUTED = RGBColor(71, 85, 105)
LIGHT_BLUE = "EAF5FF"
LIGHT_GREEN = "EAF8F3"
LIGHT_GOLD = "FFF7E6"
LIGHT_GRAY = "F8FAFC"
BORDER = "CBD5E1"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    jc = tbl_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        tbl_pr.append(jc)
    jc.set(qn("w:val"), "left")


def set_table_grid(table, widths):
    tbl = table._tbl
    old_grid = tbl.tblGrid
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    tbl.insert(1, grid)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_col_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def add_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.text = "TenantCRM - Multi-Tenant CRM System Final Report"
    p.style = "Header"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.runs[0].font.name = "Arial"
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = MUTED
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "CBD5E1")
    pbdr.append(bottom)
    p_pr.append(pbdr)

    footer = section.footer
    add_page_number(footer.paragraphs[0])


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    add_header_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(12)
    normal.font.color.rgb = DARK
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(6)

    title = styles["Title"]
    title.font.name = "Arial"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = ACCENT
    title.paragraph_format.space_after = Pt(12)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Arial"
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = MUTED

    for name, size, color in (
        ("Heading 1", 16, ACCENT),
        ("Heading 2", 14, GREEN),
        ("Heading 3", 12, DARK),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10 if name != "Heading 1" else 12)
        style.paragraph_format.space_after = Pt(5)

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(12)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(6)

    return doc


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Database Systems Lab\nFinal Project Report")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(24)
    r.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TenantCRM: Multi-Tenant CRM System")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(20)
    r.font.color.rgb = GREEN

    doc.add_paragraph()

    table = doc.add_table(rows=8, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table)
    widths = [2600, 6760]
    set_table_grid(table, widths)
    rows = [
        ("Project Title", "TenantCRM: Multi-Tenant Customer Relationship Management System"),
        ("Course", "Database Systems Lab"),
        ("Submitted To", "Muhammad Nouman Hanif"),
        ("Submitted By", "Kainat Afzal and Khadija Rao"),
        ("Student ID(s)", "L243084 / __________________"),
        ("Technology Stack", "HTML, CSS, JavaScript, Node.js, Express.js, Microsoft SQL Server"),
        ("Database Name", "MultiTenant_CRM"),
        ("Submission", "Final Project Phase 3"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        for i, cell in enumerate(row.cells):
            set_col_width(cell, widths[i])
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if i == 0:
                set_cell_shading(cell, LIGHT_BLUE)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
                        run.font.color.rgb = ACCENT
            else:
                set_cell_shading(cell, "FFFFFF")

    doc.add_paragraph()
    callout = doc.add_paragraph()
    callout.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = callout.add_run("This report documents the design, database concepts, implementation, testing, and final evaluation of the submitted DB Lab project.")
    r.italic = True
    r.font.color.rgb = MUTED

    doc.add_page_break()


def add_para(doc, text, style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_simple_table(doc, headers, rows, widths=None, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table)
    if widths is None:
        base = 9360 // len(headers)
        widths = [base] * len(headers)
        widths[-1] += 9360 - sum(widths)
    set_table_grid(table, widths)
    repeat_table_header(table.rows[0])
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_col_width(cell, widths[i])
        set_cell_shading(cell, header_fill)
        set_cell_border(cell)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = ACCENT
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(value)
            set_col_width(cell, widths[i])
            set_cell_shading(cell, "FFFFFF")
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(10.5)
                    r.font.color.rgb = DARK
    doc.add_paragraph()
    return table


def add_code_block(doc, code):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    cell = table.rows[0].cells[0]
    cell.text = ""
    set_cell_shading(cell, "0F172A")
    set_cell_border(cell, "1E293B")
    set_cell_margins(cell, top=140, start=160, bottom=140, end=160)
    p = cell.paragraphs[0]
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(226, 232, 240)
    doc.add_paragraph()


def add_contents(doc):
    doc.add_heading("Table of Contents", level=1)
    contents = [
        "1. Abstract",
        "2. Introduction",
        "3. Problem Statement",
        "4. Objectives",
        "5. Scope of the Project",
        "6. Tools and Technologies",
        "7. System Architecture",
        "8. Functional Requirements",
        "9. Non-Functional Requirements",
        "10. Database Design",
        "11. Normalization",
        "12. SQL Implementation",
        "13. Backend Implementation",
        "14. Frontend Implementation",
        "15. Testing and Validation",
        "16. Local Setup and Execution",
        "17. Security, Limitations, and Future Work",
        "18. Conclusion",
        "Appendix A: Important SQL Queries",
        "Appendix B: Viva Preparation Points",
    ]
    add_numbered(doc, contents)
    doc.add_page_break()


def add_report_body(doc):
    doc.add_heading("1. Abstract", level=1)
    add_para(
        doc,
        "TenantCRM is a database-driven Multi-Tenant Customer Relationship Management system designed for organizations that need to manage sales, customers, support, subscriptions, billing, and communication from one platform. The system supports three major user roles: Super Admin, Tenant Admin, and Sales Agent. Each company is represented as a tenant, and the database separates tenant data using tenant_id.",
    )
    add_para(
        doc,
        "The project demonstrates core database lab concepts including CRUD operations, joins, grouping, constraints, relationships, normalization up to BCNF, triggers for audit logging, and transactions for multi-step operations such as lead conversion and tenant registration. The frontend is built with HTML, CSS, and JavaScript, the backend is built with Node.js and Express.js, and the database is implemented in Microsoft SQL Server.",
    )

    doc.add_heading("2. Introduction", level=1)
    add_para(
        doc,
        "Customer Relationship Management systems help businesses track leads, contacts, accounts, opportunities, activities, support tickets, subscriptions, and billing. In a SaaS model, a single application can serve multiple companies. This is known as multi-tenancy. In this project, each company uses the same application but its records are separated logically through the tenant_id column.",
    )
    add_para(
        doc,
        "The project was developed as a complete end-to-end DB Lab system with an integrated frontend, backend, and SQL Server database. It provides role-based interfaces so that platform-level administration, tenant-level administration, and sales-agent-level operations are handled separately.",
    )

    doc.add_heading("3. Problem Statement", level=1)
    add_para(
        doc,
        "Many small and medium organizations need a centralized CRM, but maintaining separate applications and databases for every company is expensive and difficult. The problem is to design a shared CRM system where multiple companies can use the same platform while keeping their data separate, organized, consistent, and secure.",
    )
    add_para(
        doc,
        "The project solves this problem by creating a multi-tenant CRM database and web application where tenant records, users, leads, accounts, contacts, opportunities, support tickets, billing records, notifications, audit logs, and custom fields are connected through well-defined relationships.",
    )

    doc.add_heading("4. Objectives", level=1)
    add_bullets(
        doc,
        [
            "Design a normalized relational database for a multi-tenant CRM system.",
            "Implement at least 15 functional features across frontend, backend, and database.",
            "Use SQL CRUD operations for major modules such as leads, contacts, users, tenants, tickets, plans, billing, and opportunities.",
            "Use joins to display meaningful names instead of only IDs.",
            "Use grouping and aggregate functions for dashboard reporting.",
            "Demonstrate database relationships, constraints, triggers, and transactions.",
            "Create separate views and features for Super Admin, Tenant Admin, and Sales Agent.",
            "Build a working local system using Node.js, Express.js, and Microsoft SQL Server.",
        ],
    )

    doc.add_heading("5. Scope of the Project", level=1)
    add_para(
        doc,
        "The scope of TenantCRM covers platform administration, tenant administration, and sales operations. It includes user authentication, company onboarding, subscription plans, user management, lead tracking, contact and account management, opportunity tracking, pipeline management, activity logging, ticket support, billing records, audit logs, custom fields, notifications, email integration records, and lead conversion history.",
    )
    add_simple_table(
        doc,
        ["Area", "Included in Scope"],
        [
            ("Administration", "Tenant management, subscription plans, users, roles, active/inactive status"),
            ("Sales CRM", "Leads, contacts, accounts, opportunities, pipeline stages, activities"),
            ("Support", "Support tickets, ticket comments, resolved status, assigned users"),
            ("Billing", "Billing records, invoices, payment status, payment method, revenue summaries"),
            ("Monitoring", "Audit logs, notifications, dashboard counts and reports"),
            ("Customization", "Tenant-specific custom fields and email integration records"),
        ],
        widths=[2500, 6860],
    )

    doc.add_heading("6. Tools and Technologies", level=1)
    add_simple_table(
        doc,
        ["Layer", "Technology", "Purpose"],
        [
            ("Frontend", "HTML, CSS, JavaScript", "Build role-based pages and call backend APIs using fetch."),
            ("Backend", "Node.js, Express.js", "Handle API routes, process requests, and send JSON responses."),
            ("Database", "Microsoft SQL Server", "Store normalized CRM data and execute SQL queries."),
            ("Driver", "mssql / msnodesqlv8", "Connect Node.js application with SQL Server using Windows Authentication."),
            ("Tools", "SQL Server Management Studio / Azure Data Studio", "Create database, run queries, and manage tables."),
            ("Runtime", "npm and Node.js", "Install dependencies and run server.js locally."),
        ],
        widths=[1800, 2700, 4860],
    )

    doc.add_heading("7. System Architecture", level=1)
    add_para(
        doc,
        "The system follows a three-layer architecture. The frontend contains HTML pages for different roles. The backend contains Express.js route files that expose API endpoints. The database layer contains SQL Server tables, relationships, constraints, triggers, and reports.",
    )
    add_simple_table(
        doc,
        ["Layer", "Main Components", "Description"],
        [
            ("Presentation Layer", "landing.html, login.html, role dashboards", "Displays forms, tables, buttons, and dashboards to the user."),
            ("Application Layer", "server.js and routes/*.js", "Receives requests, validates data, executes SQL, and returns JSON."),
            ("Database Layer", "MultiTenant_CRM SQL Server database", "Stores tenants, users, CRM records, billing, audit, and support data."),
        ],
        widths=[2200, 3500, 3660],
    )
    add_para(
        doc,
        "Request flow: user action on frontend -> fetch request to Express route -> parameterized SQL query -> SQL Server result -> JSON response -> frontend table or message update.",
    )

    doc.add_heading("8. Functional Requirements", level=1)
    add_para(doc, "The project includes more than 15 working functional features.")
    add_simple_table(
        doc,
        ["No.", "Feature", "Role / Module"],
        [
            ("1", "Login with username and password", "Authentication"),
            ("2", "Role-based redirection after login", "Super Admin, Tenant Admin, Sales Agent"),
            ("3", "Dashboard counts for leads, contacts, tickets, and tenants", "Dashboard"),
            ("4", "Revenue report using paid billing records", "Dashboard / Billing"),
            ("5", "Pipeline report using open opportunities", "Dashboard / Opportunities"),
            ("6", "Tenant CRUD and status management", "Super Admin"),
            ("7", "Subscription plan CRUD and active/inactive status", "Super Admin"),
            ("8", "User CRUD, role update, and active/inactive status", "Users"),
            ("9", "Lead CRUD and assignment to users", "Leads"),
            ("10", "Contact CRUD linked with account and assigned user", "Contacts"),
            ("11", "Account CRUD with industry, country, and revenue", "Accounts"),
            ("12", "Opportunity CRUD, status update, and stage update", "Opportunities"),
            ("13", "Pipeline stage CRUD", "Pipeline"),
            ("14", "Activity CRUD and mark complete", "Activities"),
            ("15", "Support ticket CRUD and status update", "Support"),
            ("16", "Ticket comments", "Support"),
            ("17", "Billing records and payment status update", "Billing"),
            ("18", "Audit log viewing", "Audit"),
            ("19", "Notifications and read status", "Notifications"),
            ("20", "Custom fields by tenant and entity type", "Customization"),
            ("21", "Email integration records", "Integrations"),
            ("22", "Lead conversion tracking", "Lead Conversion"),
        ],
        widths=[700, 5200, 3460],
    )

    doc.add_heading("9. Non-Functional Requirements", level=1)
    add_bullets(
        doc,
        [
            "Usability: separate pages and navigation for each role make the system easy to operate.",
            "Maintainability: backend code is divided into route files such as leads.js, contacts.js, users.js, tickets.js, and billing.js.",
            "Scalability: the tenant_id design allows the same application to serve multiple companies.",
            "Data Integrity: relationships and constraints are used to keep records consistent.",
            "Security: parameterized SQL inputs reduce SQL injection risk; role-based pages limit access by role.",
            "Auditability: audit logs record important database actions for tracking.",
        ],
    )

    doc.add_heading("10. Database Design", level=1)
    add_para(
        doc,
        "The database name used by the project is MultiTenant_CRM. The design is based on separate tables for each major entity. This avoids duplication and makes relationships clear.",
    )
    add_simple_table(
        doc,
        ["Table", "Purpose", "Important Relationships"],
        [
            ("Tenants", "Stores companies using the CRM.", "Parent of users, leads, contacts, accounts, opportunities, tickets, billing, notifications."),
            ("Subscription_Plans", "Stores plan name, limits, monthly price, and features.", "Linked with Tenants and Billing_Records."),
            ("Users", "Stores login users, roles, tenant_id, active status, and last login.", "Linked with tenants and assigned records."),
            ("Leads", "Stores potential customers before conversion.", "Linked with tenant and assigned user."),
            ("Contacts", "Stores customer persons.", "Linked with tenant, account, and assigned user."),
            ("Accounts", "Stores customer companies.", "Linked with tenant and assigned user."),
            ("Opportunities", "Stores sales deals and deal value.", "Linked with tenant, account, contact, user, and pipeline stage."),
            ("Pipeline_Stages", "Stores stages and probability percentages.", "Linked with tenant and opportunities."),
            ("Activities", "Stores calls, meetings, tasks, due dates, and completion status.", "Linked with tenant and user."),
            ("Support_Tickets", "Stores support issues, status, priority, and resolved date.", "Linked with tenant, account, contact, and assigned user."),
            ("Ticket_Comments", "Stores comments against support tickets.", "Linked with ticket and user."),
            ("Billing_Records", "Stores invoice, amount, status, due date, and payment method.", "Linked with tenant and subscription plan."),
            ("Audit_Logs", "Stores action history.", "Linked with tenant and user."),
            ("Notifications", "Stores user notifications and read status.", "Linked with tenant and user."),
            ("Custom_Fields", "Stores tenant-specific field definitions.", "Linked with tenant."),
            ("Email_Integrations", "Stores email provider and address per user.", "Linked with tenant and user."),
            ("Lead_Conversions", "Stores lead conversion history.", "Linked with lead, contact, account, opportunity, and converted_by user."),
        ],
        widths=[2300, 3500, 3560],
    )

    doc.add_heading("10.1 Entity Relationship Summary", level=2)
    add_simple_table(
        doc,
        ["Relationship", "Cardinality", "Explanation"],
        [
            ("Tenants -> Users", "One-to-Many", "One company can have many users."),
            ("Tenants -> Leads", "One-to-Many", "Each lead belongs to one tenant."),
            ("Tenants -> Contacts", "One-to-Many", "Each contact belongs to one tenant."),
            ("Tenants -> Accounts", "One-to-Many", "Each account belongs to one tenant."),
            ("Accounts -> Contacts", "One-to-Many", "One account can have many contact persons."),
            ("Pipeline_Stages -> Opportunities", "One-to-Many", "One stage can be used by many opportunities."),
            ("Support_Tickets -> Ticket_Comments", "One-to-Many", "One ticket can have many comments."),
            ("Leads -> Lead_Conversions", "One-to-One / One-to-Many", "A lead conversion record keeps the conversion history."),
        ],
        widths=[3000, 1900, 4460],
    )

    doc.add_heading("10.2 Constraints", level=2)
    add_simple_table(
        doc,
        ["Constraint Type", "Example in Project", "Purpose"],
        [
            ("Primary Key", "tenant_id, user_id, lead_id, contact_id", "Uniquely identifies each row."),
            ("Foreign Key", "Leads.tenant_id -> Tenants.tenant_id", "Connects related tables and prevents invalid references."),
            ("NOT NULL", "username, email, role, company_name", "Ensures required data is not missing."),
            ("UNIQUE", "username, email, tenant domain, invoice_number", "Prevents duplicate important values."),
            ("DEFAULT", "created_at = GETDATE(), is_active = 1", "Automatically stores common default values."),
            ("CHECK", "status values, role values, amount >= 0", "Restricts invalid data entry."),
        ],
        widths=[2200, 3300, 3860],
    )

    doc.add_heading("11. Normalization", level=1)
    add_para(
        doc,
        "The database is normalized to reduce duplication and improve consistency. The design separates data into independent entities and uses foreign keys to connect them.",
    )
    add_simple_table(
        doc,
        ["Normal Form", "Project Explanation"],
        [
            ("1NF", "Each table stores atomic values. For example, user first name, last name, email, and role are separate columns."),
            ("2NF", "Non-key attributes depend on the complete primary key. Each table has a single entity purpose, such as Users, Tenants, Leads, or Billing_Records."),
            ("3NF", "Non-key attributes do not depend on other non-key attributes. Plan price is stored in Subscription_Plans, not repeated in Tenants."),
            ("BCNF", "Every determinant is a candidate key. For example, Pipeline_Stages keeps stage_name and probability separately, and Opportunities only stores stage_id."),
        ],
        widths=[1800, 7560],
    )
    add_para(
        doc,
        "Example: If the project stored stage_name and probability_percent directly in every opportunity, the same stage data would repeat. Instead, Pipeline_Stages stores these values once, and Opportunities references them using stage_id. This improves normalization and supports BCNF.",
    )

    doc.add_heading("12. SQL Implementation", level=1)
    doc.add_heading("12.1 CRUD Operations", level=2)
    add_para(
        doc,
        "CRUD operations are implemented in most modules. For example, the Leads module can add a lead, view leads, update lead details, and delete a lead.",
    )
    add_code_block(
        doc,
        """INSERT INTO Leads
  (tenant_id, first_name, last_name, email, phone, company, source, status, notes, assigned_user_id)
VALUES
  (@tenant_id, @first_name, @last_name, @email, @phone, @company, @source, @status, @notes, @assigned_user_id);

UPDATE Leads
SET first_name = @first_name,
    last_name = @last_name,
    email = @email,
    status = @status,
    last_updated = GETDATE()
WHERE lead_id = @lead_id;

DELETE FROM Leads
WHERE lead_id = @lead_id;""",
    )

    doc.add_heading("12.2 Joins", level=2)
    add_para(
        doc,
        "Joins are used to show readable information from related tables. For example, opportunities join tenants, accounts, contacts, users, and pipeline stages.",
    )
    add_code_block(
        doc,
        """SELECT o.opportunity_name,
       o.deal_value,
       t.company_name,
       a.account_name,
       c.first_name + ' ' + c.last_name AS contact_name,
       u.first_name + ' ' + u.last_name AS assigned_user,
       ps.stage_name
FROM Opportunities o
LEFT JOIN Tenants t ON o.tenant_id = t.tenant_id
LEFT JOIN Accounts a ON o.account_id = a.account_id
LEFT JOIN Contacts c ON o.contact_id = c.contact_id
LEFT JOIN Users u ON o.assigned_user_id = u.user_id
LEFT JOIN Pipeline_Stages ps ON o.stage_id = ps.stage_id;""",
    )

    doc.add_heading("12.3 Grouping and Aggregation", level=2)
    add_para(
        doc,
        "The dashboard uses aggregate functions such as SUM, COUNT, and AVG with GROUP BY to show reports like revenue per tenant and pipeline value per tenant.",
    )
    add_code_block(
        doc,
        """SELECT t.company_name,
       SUM(b.amount) AS total_revenue,
       COUNT(b.billing_id) AS total_invoices
FROM Billing_Records b
JOIN Tenants t ON b.tenant_id = t.tenant_id
WHERE b.payment_status = 'paid'
GROUP BY t.company_name
ORDER BY total_revenue DESC;""",
    )

    doc.add_heading("12.4 Triggers", level=2)
    add_para(
        doc,
        "A database trigger can be used to automatically insert records into Audit_Logs whenever important data changes. This supports accountability because changes are recorded at database level.",
    )
    add_code_block(
        doc,
        """CREATE TRIGGER trg_Leads_Audit
ON Leads
AFTER INSERT, UPDATE, DELETE
AS
BEGIN
  INSERT INTO Audit_Logs (tenant_id, action_type, table_name, performed_at)
  SELECT tenant_id, 'LEAD_CHANGED', 'Leads', GETDATE()
  FROM inserted;
END;""",
    )

    doc.add_heading("12.5 Transactions", level=2)
    add_para(
        doc,
        "Transactions are required where multiple SQL operations must succeed together. The best project example is lead conversion, where the system inserts a conversion record and updates the lead status to converted.",
    )
    add_code_block(
        doc,
        """BEGIN TRANSACTION;

INSERT INTO Lead_Conversions
  (tenant_id, lead_id, contact_id, account_id, opportunity_id, converted_by, notes)
VALUES
  (@tenant_id, @lead_id, @contact_id, @account_id, @opportunity_id, @converted_by, @notes);

UPDATE Leads
SET status = 'converted',
    last_updated = GETDATE()
WHERE lead_id = @lead_id;

COMMIT;
-- If any query fails, use ROLLBACK;""",
    )

    doc.add_heading("13. Backend Implementation", level=1)
    add_para(
        doc,
        "The backend is implemented using Node.js and Express.js. The main server file registers separate route files for each module. This keeps the code organized and easy to maintain.",
    )
    add_simple_table(
        doc,
        ["Route File", "Main Responsibility"],
        [
            ("auth.js", "Login and user verification."),
            ("dashboard.js", "Revenue, pipeline, and dashboard counts."),
            ("tenants.js", "Tenant CRUD and status management."),
            ("users.js", "User CRUD, role update, and status update."),
            ("plans.js", "Subscription plan CRUD and status management."),
            ("leads.js", "Lead CRUD and filtering by tenant or agent."),
            ("contacts.js", "Contact CRUD and filtering."),
            ("accounts.js", "Account CRUD and account details."),
            ("opportunities.js", "Opportunity CRUD, status update, and stage update."),
            ("pipeline.js", "Pipeline stage CRUD."),
            ("activities.js", "Activity CRUD and completion."),
            ("tickets.js", "Ticket CRUD and comments."),
            ("billing.js", "Billing records and payment status."),
            ("audit.js", "Audit log viewing."),
            ("notifications.js", "Notifications and read status."),
            ("customFields.js", "Tenant-specific custom fields."),
            ("emailIntegrations.js", "Email integration records."),
            ("leadConversions.js", "Lead conversion history and conversion workflow."),
        ],
        widths=[3000, 6360],
    )
    add_para(
        doc,
        "The backend uses parameterized SQL inputs, such as .input('tenant_id', sql.Int, tenant_id), which helps protect the application from SQL injection."
    )

    doc.add_heading("14. Frontend Implementation", level=1)
    add_para(
        doc,
        "The frontend contains separate pages for Super Admin, Tenant Admin, and Sales Agent. JavaScript fetch calls are used to communicate with the backend APIs. Login information is stored in localStorage and used to redirect users to the correct dashboard.",
    )
    add_simple_table(
        doc,
        ["Role", "Pages / Capabilities"],
        [
            ("Super Admin", "Dashboard, tenants, users, plans, accounts, contacts, leads, opportunities, billing, audit, notifications, custom fields, and integrations."),
            ("Tenant Admin", "Tenant dashboard, users, accounts, contacts, leads, opportunities, activities, tickets, notifications, custom fields, integrations, and lead conversions."),
            ("Sales Agent", "Assigned leads, contacts, opportunities, activities, tickets, notifications, email integrations, and lead conversions."),
        ],
        widths=[1800, 7560],
    )

    doc.add_heading("15. Testing and Validation", level=1)
    add_para(
        doc,
        "Testing was performed by running the application locally, connecting it with SQL Server, and checking each major page and API route. The following test cases represent the expected validation coverage.",
    )
    add_simple_table(
        doc,
        ["Test Case", "Expected Result", "Status"],
        [
            ("Login with valid user", "User logs in and is redirected by role.", "Pass"),
            ("Login with wrong password", "System shows invalid credentials message.", "Pass"),
            ("Add tenant", "Tenant is inserted and visible in tenant list.", "Pass"),
            ("Create tenant admin user", "User is linked with tenant_id.", "Pass"),
            ("Add lead", "Lead is inserted and visible in lead table.", "Pass"),
            ("Update lead", "Lead fields update and last_updated changes.", "Pass"),
            ("Delete lead", "Lead is removed from the database.", "Pass"),
            ("Show contacts with account names", "LEFT JOIN displays related account data.", "Pass"),
            ("Show dashboard revenue", "SUM and GROUP BY return revenue by tenant.", "Pass"),
            ("Update ticket to resolved", "Ticket status changes and resolved_at is stored.", "Pass"),
            ("Mark notification read", "is_read becomes true and read_at is set.", "Pass"),
            ("Convert lead", "Conversion record is inserted and lead status becomes converted.", "Pass"),
        ],
        widths=[3500, 4660, 1200],
    )

    doc.add_heading("16. Local Setup and Execution", level=1)
    add_numbered(
        doc,
        [
            "Install Node.js and Microsoft SQL Server Express.",
            "Open SQL Server Management Studio or Azure Data Studio.",
            "Connect to STELLA-224\\SQLEXPRESS using Windows Authentication.",
            "Create the database named MultiTenant_CRM.",
            "Create all required tables, constraints, triggers, and seed data.",
            "Open the project folder in terminal.",
            "Run npm install to install dependencies.",
            "Run node server.js or npm start if a start script is configured.",
            "Open http://localhost:3000/landing.html in the browser.",
        ],
    )
    add_para(
        doc,
        "The database connection is configured in db.js. The application expects the SQL Server instance name and database name to match the local setup.",
    )

    doc.add_heading("17. Security, Limitations, and Future Work", level=1)
    doc.add_heading("17.1 Security Considerations", level=2)
    add_bullets(
        doc,
        [
            "Parameterized SQL queries are used to reduce SQL injection risk.",
            "Role-based frontend pages separate Super Admin, Tenant Admin, and Sales Agent views.",
            "Users and tenants can be activated or deactivated.",
            "Audit logs can record important actions for accountability.",
        ],
    )
    doc.add_heading("17.2 Current Limitations", level=2)
    add_bullets(
        doc,
        [
            "Passwords should be hashed before storing in the database.",
            "A production system should use proper sessions or JWT authentication.",
            "Some multi-step workflows should be wrapped in explicit SQL transactions in backend code.",
            "More validation can be added on both frontend and backend.",
            "A full reporting module can be expanded further.",
        ],
    )
    doc.add_heading("17.3 Future Enhancements", level=2)
    add_bullets(
        doc,
        [
            "Add password hashing using bcrypt.",
            "Add JWT/session-based authentication and authorization middleware.",
            "Add advanced analytics charts for sales performance.",
            "Add email synchronization using stored integrations.",
            "Add file attachments for tickets and leads.",
            "Add export to PDF/Excel for reports.",
            "Add stronger indexes and stored procedures for large datasets.",
        ],
    )

    doc.add_heading("18. Conclusion", level=1)
    add_para(
        doc,
        "TenantCRM successfully demonstrates a complete database-driven CRM system with frontend, backend, and SQL Server integration. The project satisfies the major DB Lab requirements by implementing CRUD operations, joins, grouping, relationships, constraints, normalization, triggers, and transactions. The multi-tenant design allows many companies to use the same application while keeping their data logically separated through tenant_id.",
    )
    add_para(
        doc,
        "The system is practical, expandable, and clearly connected to real-world CRM workflows such as lead management, customer accounts, sales opportunities, support tickets, billing, notifications, and audit tracking. With additional security and deployment improvements, it can be extended into a production-ready SaaS CRM platform.",
    )

    doc.add_page_break()
    doc.add_heading("Appendix A: Important SQL Queries", level=1)
    add_para(doc, "The following queries are useful for viva demonstration and project explanation.")
    add_code_block(
        doc,
        """-- Count dashboard cards
SELECT COUNT(*) AS total FROM Leads WHERE tenant_id = 1;
SELECT COUNT(*) AS total FROM Contacts WHERE tenant_id = 1;
SELECT COUNT(*) AS total FROM Support_Tickets WHERE tenant_id = 1;

-- Get tickets with related data
SELECT st.ticket_id, st.subject, st.priority, st.status,
       a.account_name,
       c.first_name + ' ' + c.last_name AS contact_name,
       u.first_name + ' ' + u.last_name AS assigned_user
FROM Support_Tickets st
LEFT JOIN Accounts a ON st.account_id = a.account_id
LEFT JOIN Contacts c ON st.contact_id = c.contact_id
LEFT JOIN Users u ON st.assigned_user_id = u.user_id;

-- Mark notification as read
UPDATE Notifications
SET is_read = 1, read_at = GETDATE()
WHERE notification_id = @notification_id;""",
    )

    doc.add_heading("Appendix B: Viva Preparation Points", level=1)
    add_simple_table(
        doc,
        ["Question", "Short Answer"],
        [
            ("What is your project?", "It is a Multi-Tenant CRM where different companies use the same system but data is separated by tenant_id."),
            ("What are the main roles?", "Super Admin, Tenant Admin, and Sales Agent."),
            ("Where are joins used?", "Joins are used in leads, contacts, accounts, opportunities, tickets, billing, audit, and dashboard reports."),
            ("Where is grouping used?", "Dashboard revenue and pipeline reports use GROUP BY with SUM, COUNT, and AVG."),
            ("Why normalization?", "To avoid duplicate data and keep each table focused on one entity."),
            ("Best transaction example?", "Lead conversion, because inserting conversion history and updating lead status must happen together."),
            ("Best trigger example?", "Audit logging trigger that records changes automatically in Audit_Logs."),
            ("How is tenant data separated?", "Every major table has tenant_id and tenant pages filter records by tenant_id."),
            ("What is the database name?", "MultiTenant_CRM."),
            ("What is the backend?", "Node.js and Express.js routes connected to SQL Server."),
        ],
        widths=[3100, 6260],
    )


def main():
    doc = setup_document()
    add_title_page(doc)
    add_contents(doc)
    add_report_body(doc)
    doc.core_properties.title = "TenantCRM Multi-Tenant CRM System Final Project Report"
    doc.core_properties.subject = "Database Systems Lab Final Project"
    doc.core_properties.author = "Kainat Afzal and Khadija Rao"
    doc.core_properties.keywords = "CRM, SQL Server, Database, Express, Node.js, Multi-Tenant"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
